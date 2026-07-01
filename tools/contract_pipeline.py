#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合同批量识别管线 (unified pipeline)
从已OCR的.docx文件提取合同信息，输出标准数据结构

用法:
  python3 contract_pipeline.py scan              # 扫描合同库，列出待处理合同
  python3 contract_pipeline.py process           # 批量处理所有待处理合同
  python3 contract_pipeline.py process --force   # 强制重新处理所有合同
  python3 contract_pipeline.py status            # 查看处理状态统计

工作目录: /home/samuel/.openclaw/workspace/pm-director/tools/
"""

import os
import sys
import json
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple

# ---- 配置 ----
WORKSPACE = Path('/home/samuel/.openclaw/workspace/pm-director')
OCR_LIB = Path('/home/samuel/OCdoc/已ocr合同')
CACHE_DIR = WORKSPACE / 'cache' / 'contracts'
STANDARD_DATA_DIR = CACHE_DIR  # 标准数据JSON存放位置
DB_PATH = WORKSPACE / 'database' / 'project_management.db'

# ---- 合同编号提取模式 ----
PROJECT_ID_PATTERNS = [
    r'(ZH02-\d{9})',           # ZH02-202601001
    r'(ZH\d{2}-\d{9})',        # ZH202601001
    r'(SGSC[A-Z0-9]{14})',     # SGSC... (国网合同编号)
]

# ---- 金额提取模式 ----
AMOUNT_PATTERNS = [
    r'(\d+\.?\d*)\s*万元',
    r'(\d+\.?\d*)\s*万',
    r'(\d+)\s*元(?![\d/])',
    r'合同金额[：:]\s*人民币[（(]?([\d,]+\.?\d*)[)）]?\s*万元',
    r'金额[：:]\s*人民币[（(]?([\d,]+\.?\d*)[)）]?\s*元',
    r'¥\s*([\d,]+\.?\d*)',
]


class ChineseNumberConverter:
    """中文数字 ↔ 阿拉伯数字转换"""
    _NUM_MAP = {
        '零': 0, '一': 1, '二': 2, '两': 2, '三': 3, '四': 4,
        '五': 5, '六': 6, '七': 7, '八': 8, '九': 9
    }
    _UNIT_MAP = {'十': 10, '百': 100, '千': 1000, '万': 10000, '亿': 100000000}

    @classmethod
    def to_int(cls, s: str) -> int:
        """中文数字转int, eg '十五'→15, '一百二十三'→123"""
        if not s:
            return 0
        # Pure digits
        if s.isdigit():
            return int(s)
        s = s.strip()
        # Handle simple cases
        if s in cls._NUM_MAP:
            return cls._NUM_MAP[s]
        # Try parsing like "一", "二" etc
        result = 0
        current = 0
        for ch in s:
            if ch in cls._NUM_MAP:
                current = cls._NUM_MAP[ch]
            elif ch in cls._UNIT_MAP:
                unit = cls._UNIT_MAP[ch]
                if current == 0:
                    current = 1
                result += current * unit
                current = 0
            else:
                current = 0
        return result + current


class DocxTextExtractor:
    """从.docx提取纯文本"""

    @staticmethod
    def extract(filepath: Path) -> str:
        try:
            import docx
            doc = docx.Document(str(filepath))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        tables_text.append(' | '.join(cells))
            all_text = '\n'.join(paragraphs)
            if tables_text:
                all_text += '\n\n[表格内容]\n' + '\n'.join(tables_text)
            return all_text
        except Exception as e:
            print(f"  ⚠️  docx读取失败: {e}")
            return ""


class ContractTextParser:
    """从合同文本中提取结构化信息"""

    @staticmethod
    def extract_basic_info(text: str, filename: str) -> Dict:
        """提取基本信息"""
        info = {
            '合同编号': None,
            '财务编号': None,
            '项目名称': None,
            '甲方': None,
            '乙方': None,
            '项目负责人': None,
            '签订日期': None,
            '有效期限': None,
        }

        # 从文件名提取项目编号
        for pat in PROJECT_ID_PATTERNS:
            m = re.search(pat, filename)
            if m:
                info['合同编号'] = m.group(1)
                break

        # 从文件名提取项目名称(启发式: 取公司名后的内容)
        if not info['项目名称']:
            # 清理文件名
            name = filename.replace('.docx', '').replace('.pdf', '')
            # 去掉前导数字
            name = re.sub(r'^\d+[、.．\s]+', '', name)
            # 去掉项目编号前缀
            name = re.sub(r'^ZH\d{2}-\d{9}\s*', '', name)
            name = re.sub(r'^SGSC[A-Z0-9]+\s*', '', name)
            # 去掉尾部日期
            name = re.sub(r'\d{4}\.\d{1,2}\.\d{1,2}(\(\d\))?\s*$', '', name)
            name = re.sub(r'\(\d\)\s*$', '', name)
            name = re.sub(r'（\d）\s*$', '', name)
            name = name.strip(' -(（）)')
            info['项目名称'] = name if len(name) > 4 else None

        # 甲方识别
        patterns_a = [
            r'甲方[：:]\s*([^\n]+?)(?=乙方|\n)',
            r'委托方[（(]甲方[)）][：:]\s*([^\n]+)',
            r'发包人[（(]甲方[)）][：:]\s*([^\n]+)',
        ]
        for pat in patterns_a:
            m = re.search(pat, text[:3000])
            if m:
                info['甲方'] = m.group(1).strip()
                break

        # 乙方识别
        patterns_b = [
            r'乙方[：:]\s*([^\n]+?)(?=签订|本合同|\n)',
            r'受托方[（(]乙方[)）][：:]\s*([^\n]+)',
            r'承包人[（(]乙方[)）][：:]\s*([^\n]+)',
        ]
        for pat in patterns_b:
            m = re.search(pat, text[:3000])
            if m:
                info['乙方'] = m.group(1).strip()
                break

        # 签订日期
        date_pats = [
            r'签订日期[：:]\s*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'合同签订[日期日][：:]\s*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'签订时间[：:]\s*(\d{4}年\d{1,2}月\d{1,2}日)',
            r'本合同于\s*(\d{4}年\d{1,2}月\d{1,2}日)\s*签订',
        ]
        for pat in date_pats:
            m = re.search(pat, text[:5000])
            if m:
                info['签订日期'] = ContractTextParser._normalize_date(m.group(1))
                print(f"    📅 签订日期: {info['签订日期']}")
                break

        return info

    @staticmethod
    def _normalize_date(s: str) -> str:
        """规范化日期格式 YYYY-MM-DD"""
        m = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', s)
        if m:
            return f'{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}'
        m = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月', s)
        if m:
            return f'{m.group(1)}-{int(m.group(2)):02d}'
        m = re.search(r'(\d{4})-(\d{1,2})-(\d{1,2})', s)
        if m:
            return f'{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}'
        return s

    @staticmethod
    def detect_project_type(text: str, project_name: str) -> str:
        """检测项目类型: research / service"""
        research_kw = ['研究', '研发', '课题', '关键技术', '科技项目', '技术创新']
        service_kw = ['服务', '运维', '维护', '支撑', '实施', '建设', '测绘', '监测', '咨询']

        r_score = sum(2 for kw in research_kw if kw in project_name) + \
                  sum(1 for kw in research_kw if kw in text[:5000])
        s_score = sum(2 for kw in service_kw if kw in project_name) + \
                  sum(1 for kw in service_kw if kw in text[:5000])

        if r_score >= s_score and r_score >= 3:
            return '科研类'
        elif s_score > r_score and s_score >= 3:
            return '服务类'

        # fallback: check chapter structure
        if any(kw in text[:8000] for kw in ['研究内容', '研究计划', '项目进度安排']):
            return '科研类'
        if any(kw in text[:8000] for kw in ['服务内容', '服务要求', '服务期限', '工期']):
            return '服务类'

        return '未知'

    @staticmethod
    def extract_amount(text: str) -> Optional[Dict]:
        """提取合同金额"""
        text_head = text[:8000]
        for pat in AMOUNT_PATTERNS:
            m = re.search(pat, text_head)
            if m:
                val_str = m.group(1).replace(',', '')
                try:
                    val = float(val_str)
                    if '万元' in m.group(0) or '万' in m.group(0):
                        return {'含税金额': int(val * 10000), '不含税金额': None, '税率': None, '税额': None}
                    else:
                        return {'含税金额': int(val), '不含税金额': None, '税率': None, '税额': None}
                except:
                    continue
        return None

    @staticmethod
    def extract_research_stages(text: str) -> List[Dict]:
        """提取研究阶段 (for research contracts)"""
        stages = []

        # Strategy 1: Look for structured "阶段X" patterns
        # Pattern: 阶段编号 + 时间范围 + 主要内容 + 考核目标
        stage_blocks = re.split(r'\n\s*(?=\d+\.\d+\s+|第[一二三四五六七八九十\d]+阶段|\d+[、.．]\s*[阶][段][名称]?\s*)', text)

        for block in stage_blocks:
            if len(block) < 50:
                continue

            # Try to extract stage number
            stage_num = None
            m = re.search(r'(?:阶段|第)\s*[一二三四五六七八九十\d]+\s*[阶段条]', block[:100])
            if m:
                num_text = re.search(r'[一二三四五六七八九十\d]+', m.group())
                if num_text:
                    stage_num = num_text.group()
                    if not stage_num.isdigit():
                        stage_num = str(ChineseNumberConverter.to_int(stage_num))

            if not stage_num:
                m = re.search(r'(\d+)\.\s', block[:100])
                if m:
                    stage_num = m.group(1)

            if not stage_num:
                continue

            # Extract time range
            time_range = ''
            tm = re.search(r'(\d{4}年\d{0,2}月?\d{0,2}日?)\s*[至到~—–-]\s*(\d{4}年?\d{0,2}月?\d{0,2}日?)', block)
            if tm:
                time_range = f'{tm.group(1).strip()} 至 {tm.group(2).strip()}'

            # Extract content and assessment
            lines = block.split('\n')
            content_lines = []
            assessment_lines = []

            in_content = False
            in_assessment = False
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if re.search(r'研究内容[：:]|主要内容[：:]', line):
                    in_content = True
                    in_assessment = False
                    continue
                if re.search(r'考核目标[：:]|考核[：:]', line):
                    in_assessment = True
                    in_content = False
                    content = ' '.join(content_lines).strip()
                    continue
                if re.search(r'阶段名称[：:]|时间[：:]|时间范围[：:]', line):
                    continue
                if in_content:
                    content_lines.append(line)
                elif in_assessment:
                    assessment_lines.append(line)

            content = ' '.join(content_lines).strip()
            assessment = ' '.join(assessment_lines).strip()

            if not content and not assessment:
                # Fallback: try paragraph after stage header
                idx = block.find('\n')
                if idx > 0:
                    rest = block[idx:].strip()
                    content = rest[:min(200, len(rest))].split('\n')[0]

            stage = {
                '阶段编号': stage_num,
                '时间范围': time_range,
                '主要内容': content if content else '',
                '考核目标': assessment if assessment else '',
            }
            stages.append(stage)

        # Strategy 2: Table-style extraction (OCR pipes format)
        if len(stages) < 2:
            stages = ContractTextParser._table_extract_stages(text)

        return stages

    @staticmethod
    def _table_extract_stages(text: str) -> List[Dict]:
        """从表格风格文本提取阶段"""
        stages = []
        # Look for pipe-separated table content
        table_sections = re.findall(r'\|[^|]+\|[^|]+\|[^|]+\|', text)
        if table_sections:
            for line in table_sections:
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if len(cells) >= 3:
                    stage = {
                        '阶段编号': str(len(stages) + 1),
                        '时间范围': cells[0],
                        '主要内容': cells[1] if len(cells) > 1 else '',
                        '考核目标': cells[2] if len(cells) > 2 else '',
                    }
                    stages.append(stage)
        return stages

    @staticmethod
    def extract_payment_plans(text: str) -> List[Dict]:
        """提取付款计划"""
        payments = []

        # Strategy 1: Structured payment entries
        # Pattern: 年度 + 完成XX + 金额
        payment_patterns = [
            r'(\d{4})\s*年\s*(完成[^0-9\n]{5,80}?)\s*(\d+\.?\d*)\s*万元',
            r'(第[一二三四五六七八九十\d]+[笔期][付])[：:。]?\s*(完成[^\d\n]{5,80}?)[，,]\s*(?:甲方)?[^\d]*?(\d+\.?\d*)\s*万元',
            r'付款节点[：:]?\s*(\d+)\s*[．.、]?\s*(.+?)[，,]\s*金额[：:]?\s*(\d+\.?\d*)\s*万元',
        ]

        for pat in payment_patterns:
            matches = re.findall(pat, text)
            if matches:
                for m in matches:
                    payment = {
                        '付款节点': len(payments) + 1,
                        '付款年度': m[0] if re.match(r'\d{4}', str(m[0])) else '',
                        '付款条件': '',
                        '金额': None,
                        '付款时间': '',
                    }
                    # Find the condition and amount in the match groups
                    for g in m[1:]:
                        if '完成' in str(g) or '验收' in str(g) or '签订' in str(g):
                            payment['付款条件'] = g.strip()
                        elif re.match(r'\d+\.?\d*', str(g)):
                            try:
                                val = float(str(g).replace(',', ''))
                                payment['金额'] = val * 10000  # convert to 元
                            except:
                                pass
                    payments.append(payment)

        # Strategy 2: Look for "合同签订后" or milestone patterns in service contracts
        if not payments:
            service_pat = [
                r'(合同签订[^，,。]{5,50}?)[，,]\s*(?:甲方[^，,。]{0,20}?)?支付\s*(\d+\.?\d*)%\s*[，,。]?\s*(?:即[^，,。]*?)?(\d+\.?\d*)\s*万元',
                r'(验收[^，,。]{5,50}?)[，,]\s*(?:甲方[^，,。]{0,20}?)?支付\s*(\d+\.?\d*)%\s*[，,。]?\s*(?:即[^，,。]*?)?(\d+\.?\d*)\s*万元',
                r'(质保[^，,。]{5,50}?)[，,]\s*(?:甲方[^，,。]{0,20}?)?支付\s*(\d+\.?\d*)%\s*[，,。]?\s*(?:即[^，,。]*?)?(\d+\.?\d*)\s*万元',
            ]
            for pat in service_pat:
                matches = re.findall(pat, text)
                if matches:
                    for m in matches:
                        payment = {
                            '付款节点': len(payments) + 1,
                            '付款条件': m[0].strip(),
                            '付款比例': float(m[1]),
                            '金额': float(m[2]) * 10000,
                            '付款时间': '',
                        }
                        payments.append(payment)

        return payments

    @staticmethod
    def extract_deliverables(text: str) -> List[Dict]:
        """提取交付物"""
        deliverables = []
        # Look for deliverable-related keywords
        lines = text.split('\n')
        in_deliverable = False
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if re.search(r'交付[物成果]|提交[成果]|最终成果', line[:60]):
                in_deliverable = True
                continue
            if in_deliverable and len(line) > 5 and len(line) < 100:
                # Check if this looks like a deliverable item
                if not re.search(r'^[第\d章节条]', line) and not re.match(r'^\d+[.．、]', line) if len(line) < 80 else True:
                    deliverables.append({
                        '类型': '成果',
                        '名称': line.strip('（()）'),
                        '数量': 1,
                    })
                if len(deliverables) > 5:
                    break
            if in_deliverable and re.search(r'验收|付款|其它|其他', line[:30]):
                break

        return deliverables

    @staticmethod
    def extract_service_info(text: str) -> Dict:
        """提取服务类项目信息"""
        info = {}
        # 服务期限
        m = re.search(r'(?:服务期限|合同期限|工期)[：:]\s*([^\n]+)', text[:5000])
        if m:
            info['服务期限'] = m.group(1).strip()

        # 服务内容
        content_start = re.search(r'服务内容[：:]\s*', text)
        if content_start:
            start = content_start.end()
            end = re.search(r'\n\s*\n|\n[第\d]', text[start:start+2000])
            if end:
                info['服务内容'] = text[start:start+end.start()].strip()[:500]
            else:
                info['服务内容'] = text[start:start+1000].strip()[:500]

        return info


class StandardDataBuilder:
    """构建标准数据JSON"""

    @staticmethod
    def build(text: str, filename: str) -> Dict:
        """从文本构建标准数据"""
        parser = ContractTextParser()

        # 1. Basic info
        basic = parser.extract_basic_info(text, filename)

        # 2. Project type
        ptype = parser.detect_project_type(text, basic.get('项目名称', ''))

        # 3. Amount
        amount = parser.extract_amount(text)

        # 4. Research stages or service info
        stages = []
        payments = []
        deliverables = []
        service_info = {}

        if ptype == '科研类':
            stages = parser.extract_research_stages(text)
            payments = parser.extract_payment_plans(text)
        elif ptype == '服务类':
            service_info = parser.extract_service_info(text)
            payments = parser.extract_payment_plans(text)

        deliverables = parser.extract_deliverables(text)

        # 5. Build result
        result = {
            '基本信息': basic,
            '项目类型': ptype,
            '合同金额': amount or {},
            '识别元数据': {
                '识别时间': datetime.now().strftime('%Y-%m-%dT%H:%M:%S+08:00'),
                '识别工具': 'contract_pipeline.py (auto-extraction)',
                '识别状态': '自动提取',
                '备注': f'从OCR合同库自动提取, 文本长度{len(text)}字'
            },
            '_raw_text_length': len(text),
        }

        if stages:
            result['研究阶段'] = stages
        elif service_info:
            result['服务信息'] = service_info

        if payments:
            result['付款计划'] = payments

        if deliverables:
            result['交付物'] = deliverables

        # 验收信息 (try to extract)
        acceptance_info = StandardDataBuilder._extract_acceptance(text)
        if acceptance_info:
            result['验收信息'] = acceptance_info

        return result

    @staticmethod
    def _extract_acceptance(text: str) -> Optional[Dict]:
        """提取验收信息"""
        info = {}
        m = re.search(r'验收[方方式标准][：:]?\s*([^\n]{5,100})', text[:8000])
        if m:
            info['验收方式'] = m.group(1).strip()

        m = re.search(r'验收标准[：:]?\s*([^\n]{5,100})', text[:8000])
        if m:
            info['验收标准'] = m.group(1).strip()

        return info if info else None


class ContractPipeline:
    """合同批量识别管线"""

    def __init__(self):
        self.stats = {'total': 0, 'already_done': 0, 'processed': 0, 'failed': 0, 'skipped': 0}

    def scan_ocr_library(self) -> List[Dict]:
        """扫描OCR合同库，返回所有合同文件元数据"""
        contracts = []
        if not OCR_LIB.exists():
            print(f"❌ OCR合同库路径不存在: {OCR_LIB}")
            return contracts

        for f in sorted(OCR_LIB.glob('*.docx')):
            contracts.append({
                'filepath': f,
                'filename': f.name,
                'size_kb': f.stat().st_size / 1024,
            })
        self.stats['total'] = len(contracts)
        return contracts

    def load_existing_standard_data(self) -> Dict[str, Dict]:
        """加载现有的标准数据"""
        existing = {}
        for f in STANDARD_DATA_DIR.glob('*标准数据.json'):
            try:
                data = json.load(open(f))
                name = data.get('基本信息', {}).get('项目名称', '') or data.get('project_name', '')
                status = data.get('识别元数据', {}).get('识别状态', '')
                existing[f.stem.replace('-标准数据', '')] = {
                    'data': data,
                    'status': status,
                    'file': f
                }
            except:
                pass
        return existing

    def need_process(self, contract: Dict, existing: Dict) -> Tuple[bool, str]:
        """判断是否需要处理"""
        filename = contract['filename']
        # Check if any existing standard data matches this file
        for key, edata in existing.items():
            if filename[:20] in str(edata['data']) or key in filename:
                if edata['status'] in ['完整', '手动提取', '自动提取']:
                    return False, f"已有标准数据 ({edata['status']})"
                else:
                    return True, f"已有数据但状态: {edata['status']}"
        return True, "无标准数据"

    def process_single(self, contract: Dict, force: bool = False) -> Optional[Dict]:
        """处理单个合同"""
        filepath = contract['filepath']
        filename = contract['filename']
        print(f"\n  📄 处理: {filename}")

        # Extract text
        text = DocxTextExtractor.extract(filepath)
        if not text or len(text) < 100:
            print(f"  ⚠️  文本过短({len(text) if text else 0}字), 跳过")
            return None

        print(f"  📝 文本长度: {len(text)}字")

        # Build standard data
        result = StandardDataBuilder.build(text, filename)

        # Try to fill missing fields from DB if available
        self._enrich_from_db(result)

        return result

    def _enrich_from_db(self, result: Dict):
        """从数据库补充信息"""
        try:
            import sqlite3
            conn = sqlite3.connect(str(DB_PATH))
            c = conn.cursor()
            project_name = result.get('基本信息', {}).get('项目名称', '')
            if project_name:
                # Look up in contracts table
                rows = c.execute(
                    'SELECT contract_id, project_type, contract_amount, party_a, party_b, project_leader, sign_date, service_content '
                    'FROM contracts WHERE project_name LIKE ?',
                    ('%' + project_name[:20] + '%',)
                ).fetchall()
                if rows:
                    r = rows[0]
                    basic = result.get('基本信息', {})
                    if not basic.get('项目名称') and r[1]:
                        basic['项目名称'] = r[1]
                    if not basic.get('甲方') and r[4]:
                        basic['甲方'] = r[4]
                    if not basic.get('乙方') and r[5]:
                        basic['乙方'] = r[5]
                    if not basic.get('项目负责人') and r[6]:
                        basic['项目负责人'] = r[6]
                    if not basic.get('合同编号') and r[0]:
                        basic['合同编号'] = r[0]
                    if r[3] and not result.get('合同金额', {}).get('含税金额'):
                        result['合同金额'] = {'含税金额': int(r[3] * 10000), '不含税金额': None, '税率': None, '税额': None}
            conn.close()
        except Exception as e:
            pass  # DB enrich is best-effort

    def save_result(self, filename: str, result: Dict):
        """保存标准数据JSON"""
        # Generate output filename
        contract_id = result.get('基本信息', {}).get('合同编号', '')
        if contract_id:
            out_name = f'{contract_id}-标准数据.json'
        else:
            # Use base name without extension
            base = Path(filename).stem
            out_name = f'{base}-标准数据.json'

        out_path = STANDARD_DATA_DIR / out_name
        with open(out_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"  💾 已保存: {out_name}")
        return out_path

    def run(self, force: bool = False):
        """运行管线"""
        print('=' * 80)
        print('合同批量识别管线 v1.0')
        print(f'开始时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        print('=' * 80)

        # Scan
        contracts = self.scan_ocr_library()
        existing = self.load_existing_standard_data()

        print(f'\n📦 OCR合同库: {len(contracts)} 个文件')
        print(f'📋 现有标准数据: {len(existing)} 个')

        # Process
        for i, ct in enumerate(contracts, 1):
            print(f'\n[{i}/{len(contracts)}]', end='')
            need, reason = self.need_process(ct, existing)
            if not need and not force:
                print(f'  ⏭️  {reason}')
                self.stats['skipped'] += 1
                continue

            result = self.process_single(ct, force)
            if result:
                self.save_result(ct['filename'], result)
                self.stats['processed'] += 1
            else:
                self.stats['failed'] += 1

        # Summary
        print('\n' + '=' * 80)
        print('处理完成统计')
        print(f'  总文件: {self.stats["total"]}')
        print(f'  已跳过(已有完整数据): {self.stats["skipped"]}')
        print(f'  本次处理: {self.stats["processed"]}')
        print(f'  失败: {self.stats["failed"]}')
        print('=' * 80)

    def status(self):
        """显示处理状态"""
        contracts = self.scan_ocr_library()
        existing = self.load_existing_standard_data()

        print(f'\n📦 合同库文件: {len(contracts)}')
        print(f'📋 标准数据文件: {len(existing)}')
        print()

        print(f'{"序号":>4} | {"文件":<60} | {"状态":<20} | {"类型":<10}')
        print('-' * 100)
        for i, ct in enumerate(contracts, 1):
            need, reason = self.need_process(ct, existing)
            status = '✅ 已完成' if not need else '⏳ 待处理'
            ptype = ''
            # Try to get type from existing data
            for key, edata in existing.items():
                if key in ct['filename'] or ct['filename'][:20] in str(edata['data']):
                    ptype = edata['data'].get('项目类型', '')
                    break
            print(f'{i:>4} | {ct["filename"]:<60} | {status:<20} | {ptype:<10}')


def main():
    if len(sys.argv) < 2:
        print('用法:')
        print('  python3 contract_pipeline.py scan        # 扫描所有文件状态')
        print('  python3 contract_pipeline.py process     # 批量处理')
        print('  python3 contract_pipeline.py process --force  # 强制重新处理')
        print('  python3 contract_pipeline.py status      # 详细状态')
        return

    pipeline = ContractPipeline()
    cmd = sys.argv[1]

    if cmd == 'scan':
        contracts = pipeline.scan_ocr_library()
        print(f'\n📦 OCR合同库: {len(contracts)} 个文件')
        for c in contracts:
            print(f'  {c["filename"]}')
    elif cmd == 'process':
        force = '--force' in sys.argv
        pipeline.run(force=force)
    elif cmd == 'status':
        pipeline.status()
    else:
        print(f'未知命令: {cmd}')


if __name__ == '__main__':
    main()
