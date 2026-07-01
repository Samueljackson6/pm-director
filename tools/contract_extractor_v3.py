#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合同信息提取器 V3 - 实用化版本
目标：可靠提取基本字段 + 尽力提取阶段/付款/交付物

核心策略：
  - 基本字段（项目编号、甲乙双方、金额、类型）：高准确率提取
  - 研究阶段：依赖章节定位（项目进度安排/研究内容）
  - 付款计划：依赖章节定位（项目支付计划/经费支付）
  - 交付物：关键词匹配 + 章节定位
  - 对于无法提取详细数据的合同，保留原始文本供后续处理

用法:
  python3 contract_extractor_v3.py <合同.docx>     # 提取单个合同
  python3 contract_extractor_v3.py --batch <目录>   # 批量提取
"""

import sys, os, json, re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ---- 路径配置 ----
WORKSPACE = Path('/home/samuel/.openclaw/workspace/pm-director')
OCR_LIB = Path('/home/samuel/OCdoc/已ocr合同')
CACHE_DIR = WORKSPACE / 'cache' / 'contracts'


# ============================================================
#  文本处理工具
# ============================================================
def clean(t: str) -> str:
    t = re.sub(r'\s+', ' ', t)
    return t.strip()

def norm_date(s: str) -> str:
    m = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', s)
    if m: return f'{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}'
    m = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月', s)
    if m: return f'{m.group(1)}-{int(m.group(2)):02d}'
    m = re.search(r'(\d{4})[.年](\d{1,2})[.月](\d{1,2})', s)
    if m: return f'{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}'
    return s

def parse_amount(text: str) -> Optional[float]:
    """提取金额(元). 返回None表示未找到"""
    head = text[:12000]
    pats = [
        (r'合同金额[：:]\s*[人民币￥¥]?\s*([\d,]+\.?\d*)\s*万元', 10000),
        (r'合同价款[：:]\s*[人民币￥¥]?\s*([\d,]+\.?\d*)\s*万元', 10000),
        (r'合同总价[：:]\s*[人民币￥¥]?\s*([\d,]+\.?\d*)\s*万元', 10000),
        (r'项目经费总额[^¥]*¥\s*([\d,]+\.?\d*)', 1),
        (r'经费总额[^\\d]*?([\d,]+\.?\d*)\s*万元', 10000),
        (r'人民币[（(]\s*([\d,]+\.?\d*)\s*万元[)）]', 10000),
        (r'金额[：:]\s*[人民币￥¥]?\s*([\d,]+\.?\d*)\s*万元', 10000),
        (r'总价[：:]\s*[人民币￥¥]?\s*([\d,]+\.?\d*)\s*万元', 10000),
        (r'¥\s*([\d,]+\.?\d*)\s*[),，]', 1),
    ]
    for pat, mul in pats:
        m = re.search(pat, head)
        if m:
            try:
                return float(m.group(1).replace(',', '')) * mul
            except:
                pass
    return None


# ============================================================
#  章节提取（核心）
# ============================================================
def find_section(text: str, headers: list, max_n: int = 6000) -> str:
    """在全文text中定位headers中的章节标题并返回后续内容(最多max_n字符)"""
    for h in headers:
        escaped = re.escape(h)
        patterns = [
            rf'(?:^|\n)\s*{escaped}\s*[：:]\s*\n(.{{1,{max_n}}}?)(?=\n\s*(?:第|\d+[.、]|附件|以下|\Z))',
            rf'(?:^|\n)\s*\d+[.、．]\s*{escaped}\s*\n(.{{1,{max_n}}}?)(?=\n\s*\d+[.、．]|\n\s*第|\Z)',
            rf'(?:^|\n)\s*{escaped}.*?\n(.{{1,{max_n}}}?)(?=\n\s*(?:第|\d+[.、]|附件|以下|\Z))',
        ]
        for pat in patterns:
            mch = re.search(pat, text, re.DOTALL)
            if mch:
                return mch.group(1).strip()
    return ''


# ============================================================
#  基本信息提取
# ============================================================
def extract_basic(text: str, filename: str) -> dict:
    """提取基础信息: 项目编号, 名称, 甲乙双方, 金额, 类型, 日期"""
    basic = {
        '合同编号': '', '财务编号': None, '项目名称': '',
        '甲方': '', '乙方': '', '项目负责人': None,
        '签订日期': None, '有效期限': None,
        '合同金额': None, '项目类型': '',
    }

    # --- 项目编号 ---
    # From filename
    for pat in [r'(ZH02-\d{9})', r'(SGSC[A-Z0-9]{14})', r'(SGTYHT/\d+[\w-]+)', r'(ZHH?\d{9})']:
        m = re.search(pat, filename)
        if m: basic['合同编号'] = m.group(1); break
    # From text
    if not basic['合同编号']:
        for pat in [r'(ZH02-\d{9})', r'(SGSC[A-Z0-9]{14})']:
            m = re.search(pat, text[:2000])
            if m: basic['合同编号'] = m.group(1); break

    # --- 项目名称 ---
    # From filename (clean)
    name = Path(filename).stem
    name = re.sub(r'^\d+[.、．\s]+', '', name)
    name = re.sub(r'\d{4}\.\d{1,2}\.\d{1,2}(?:[（(]\d[)）])?\s*$', '', name)
    name = re.sub(r'[（(]\d[)）]$', '', name)
    name = re.sub(r'^（已压缩）', '', name)
    name = name.strip(' -—（()）•·')
    basic['项目名称'] = name if len(name) > 6 else ''

    # --- 甲方(委托方) ---
    for pat in [r'(?:委托方|甲方)[（(]甲方[)）]?[：:]\s*([^\n]{4,60})',
                r'甲方[：:]\s*([^\n]{4,60})',
                r'发包人[：:]\s*([^\n]{4,60})']:
        m = re.search(pat, text[:3000])
        if m:
            a = clean(m.group(1))
            a = re.sub(r'\s+', '', a)
            basic['甲方'] = a
            break

    # --- 乙方(受托方) ---
    for pat in [r'(?:受托方|乙方)[（(]乙方[)）]?[：:]\s*([^\n]{4,60})',
                r'乙方[：:]\s*([^\n]{4,60})',
                r'承包人[：:]\s*([^\n]{4,60})']:
        m = re.search(pat, text[:3000])
        if m:
            b = clean(m.group(1))
            b = re.sub(r'\s+', '', b)
            basic['乙方'] = b
            break

    # --- 签订日期 ---
    for pat in [r'签订日期[：:]\s*(\d{4}年?\d{0,2}月?\d{0,2}日?)',
                r'合同签订[日期日][：:]\s*(\d{4}年?\d{0,2}月?\d{0,2}日?)',
                r'签订时间[：:]\s*(\d{4}年?\d{0,2}月?\d{0,2}日?)']:
        m = re.search(pat, text[:5000])
        if m:
            basic['签订日期'] = norm_date(m.group(1))
            break

    # --- 有效期限 ---
    for pat in [r'(?:有效期限|合同期限|合同有效期)[：:]\s*([^\n]{5,80})']:
        m = re.search(pat, text[:8000])
        if m:
            basic['有效期限'] = clean(m.group(1))
            break

    # --- 项目类型 ---
    pname = basic['项目名称']
    r_score = sum(3 for kw in ['研究', '研发', '课题', '科技项目', '关键技术'] if kw in pname)
    s_score = sum(3 for kw in ['服务', '运维', '维护', '支撑', '建设', '测绘', '监测', '咨询', '改造'] if kw in pname)
    r_score += sum(2 for kw in ['研究内容', '考核目标', '研究计划', '科技项目', '研究开发'] if kw in text[:6000])
    s_score += sum(2 for kw in ['服务内容', '服务要求', '工期', '服务期限'] if kw in text[:6000])
    basic['项目类型'] = '科研类' if r_score >= s_score and r_score >= 2 else '服务类' if s_score >= 2 else '未知'

    # --- 合同金额 ---
    amt = parse_amount(text)
    if amt:
        basic['合同金额'] = {'含税金额': int(amt), '不含税金额': None, '税率': None, '税额': None}
    else:
        basic['合同金额'] = None

    return basic


# ============================================================
#  研究阶段提取
# ============================================================
def extract_stages(text: str) -> List[Dict]:
    """提取研究阶段列表"""
    stages = []

    # Target sections
    sec = find_section(text, [
        '项目进度安排', '研究计划', '实施计划', '研究内容及进度安排',
        '进度安排', '研究内容', '主要研究内容'
    ])
    if not sec:
        return stages  # No stage section found

    # Strategy 1: Numbered stages like "5.1 XXX", "5.2 XXX"
    # These are clauses within the section
    lines = sec.split('\n')
    current = None

    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Match stage number: "5.1", "5.2", "1", "2" etc
        m_stage = re.match(r'(\d+(?:\.\d+)?)\s*[.、．]\s*(.{5,200}?)$', line)
        if m_stage:
            # Save previous
            if current and (current['主要内容'] or current['考核目标']):
                stages.append(current)
            num = m_stage.group(1)
            # Check if this is a stage (starts with 5.x for scientific contracts)
            desc = m_stage.group(2).strip()
            current = {
                '阶段编号': num,
                '时间范围': '',
                '主要内容': desc,
                '考核目标': '',
            }
            continue

        if not current:
            continue

        # Check for time range in current line
        tm = re.search(r'(\d{4}年?\d{0,2}月?\d{0,2}日?)\s*[至到~—–-]\s*(\d{4}年?\d{0,2}月?\d{0,2}日?)', line)
        if tm and not current['时间范围']:
            current['时间范围'] = f'{tm.group(1).strip()} 至 {tm.group(2).strip()}'
        elif re.search(r'考核目标[：:]', line):
            current['考核目标'] = line.split('：', 1)[-1].strip() if '：' in line else line
        elif len(line) > 10 and not re.match(r'^[第\d]', line):
            # Continuation of current stage content
            current['主要内容'] += ' ' + line

    if current and (current['主要内容'] or current['考核目标']):
        stages.append(current)

    # Clean up - remove non-stage items (contract clauses vs actual stages)
    # Keep only items that look like research stages
    real_stages = []
    for s in stages:
        content = s.get('主要内容', '') + s.get('考核目标', '')
        # Skip pure legal/ administrative clauses
        if re.search(r'风险承担|保密|陈述与保证|研究成果的归属|违约责任|合同.*变更|合同.*解除|争议解决|合同生效', content[:100]):
            continue
        # Keep if it mentions research keywords
        if any(kw in content for kw in ['研究', '开发', '设计', '分析', '测试', '试验', '算法', '模型', '系统', '专利', '论文', '报告', '数据', '方案', '验收']):
            real_stages.append(s)
        elif len(content) > 40:
            real_stages.append(s)

    return real_stages


# ============================================================
#  付款计划提取
# ============================================================
def extract_payments(text: str) -> List[Dict]:
    """提取付款计划"""
    payments = []

    # Find payment section
    sec = find_section(text, [
        '项目支付计划', '经费支付', '支付计划', '付款安排',
        '经费预算', '费用及支付', '费用与支付', '付款计划'
    ])

    # Also check "特别约定" section for payment-related clauses
    special_sec = find_section(text, ['特别约定', '其他需要补充约定的内容'])
    payment_clauses = ''
    if special_sec:
        for line in special_sec.split('\n'):
            if any(kw in line for kw in ['支付', '付款', '经费', '预留']):
                payment_clauses += line + '\n'

    search_text = (sec or '') + '\n' + payment_clauses
    if not search_text.strip():
        search_text = text  # fallback

    # Pattern 1: Year + condition + amount (科研类标准)
    for m in re.finditer(
        r'(\d{4})\s*年\s+(?:甲方提供经费[：:])?(完成[^，。\n]{5,100}?)\s*甲方[^，。\n]{0,40}?(\d+\.?\d*)\s*[万元元]',
        search_text
    ):
        amt = float(m.group(3)) * 10000 if '万' in m.group(0) else float(m.group(3))
        payments.append({
            '付款节点': len(payments) + 1,
            '付款条件': clean(m.group(2)),
            '金额': int(amt) if amt == int(amt) else amt,
            '付款时间': f"{m.group(1)}年",
        })

    # Pattern 2: 完成...第X条... (国网格式)
    if not payments:
        for m in re.finditer(
            r'(完成本合同条款[\d.~～]+?[^。\n]{5,100}?)[。，\n]\s*(\d+\.?\d*)\s*[万元元]',
            search_text
        ):
            amt = float(m.group(2)) * 10000 if '万' in m.group(0) else float(m.group(2))
            payments.append({
                '付款节点': len(payments) + 1,
                '付款条件': clean(m.group(1)),
                '金额': int(amt),
                '付款时间': '',
            })

    # Pattern 3: Service contract - 竣工结算/分期
    if not payments:
        for m in re.finditer(
            r'(竣工结算[^。\n]{10,100}?)。?\s*(?:扣除质保金后支付合同金额的(\d+)%)?\s*[^，。]*?(\d+\.?\d*)\s*[万元元]',
            search_text
        ):
            amt = float(m.group(3)) * 10000 if '万' in m.group(0) else float(m.group(3))
            payments.append({
                '付款节点': len(payments) + 1,
                '付款条件': clean(m.group(1)),
                '付款比例': int(m.group(2)) if m.group(2) else 100,
                '金额': int(amt),
            })

    # Pattern 4: 分期支付
    if not payments:
        for m in re.finditer(
            r'(?:第[一二三四五六七八九十\d]+[期笔次])(?:付款?)[：:，,。]?\s*(完成[^。\n]{5,120}?)[。，,]?\s*(?:甲方[^。\n]{0,40}?)?[支付]?\s*(\d+\.?\d*)\s*[万元元]',
            search_text
        ):
            amt = float(m.group(2)) * 10000 if '万' in m.group(0) else float(m.group(2))
            payments.append({
                '付款节点': len(payments) + 1,
                '付款条件': clean(m.group(1)),
                '金额': int(amt),
            })

    # Pattern 5: 支付比例 from special clauses
    if not payments:
        for m in re.finditer(
            r'(\d{4})年\s*第(\d+)\s*笔\s*(?:支付款项)?\s*(.*?)(?:为预留经费|支付)',
            payment_clauses
        ):
            payments.append({
                '付款节点': int(m.group(2)),
                '付款条件': clean(m.group(3)) if m.group(3) else f'{m.group(1)}年第{m.group(2)}笔',
                '金额': 0,  # Unknown
                '付款时间': f'{m.group(1)}年',
            })

    return payments


# ============================================================
#  交付物提取
# ============================================================
def extract_deliverables(text: str) -> List[Dict]:
    """提取交付物"""
    deliverables = []

    # Look in 最终成果  section
    sec = find_section(text, ['最终成果形式', '成果形式', '交付物', '交付成果', '技术成果', '成果清单'])
    if sec:
        for line in sec.split('\n'):
            line = line.strip()
            if not line or len(line) < 3:
                continue
            if re.match(r'^[第\d]', line):
                continue
            # Check for deliverable pattern
            for m in re.finditer(r'(申请\s*发明\s*专利\s*\d*\s*项)|(软件著作权\s*\d*\s*项)|(论文\s*\d*\s*篇)|(技术报告)|(系统\s*\d*\s*套)|(模型\s*\d*\s*套)|(装置\s*\d*\s*套)|(算法\s*\d*\s*套)', line):
                deliverables.append({'类型': '成果', '名称': clean(m.group()), '数量': 1})

    # Fallback: full-text keyword search
    if not deliverables:
        kw_pairs = [
            ('专利', r'(?:申请|申报)?\s*(?:发明)?\s*专利\s*\d*\s*项'),
            ('论文', r'(?:发表|撰写)?\s*论文\s*\d*\s*篇'),
            ('软件著作权', r'(?:软件)?\s*著作权\s*\d*\s*项'),
            ('报告', r'(?:技术|研究|调研|验收)?\s*报告\s*\d*\s*[份篇]'),
            ('系统', r'(?:软件|信息)?\s*系统\s*\d*\s*套'),
            ('模型', r'(?:预测|分析|仿真)?\s*模型\s*\d*\s*套'),
        ]
        seen = set()
        for dtype, pat in kw_pairs:
            for m in re.finditer(pat, text[:15000]):
                name = clean(m.group())
                if name not in seen and len(name) > 3:
                    seen.add(name)
                    deliverables.append({'类型': dtype, '名称': name, '数量': 1})

    return deliverables


# ============================================================
#  服务信息提取
# ============================================================
def extract_service(text: str) -> Dict:
    info = {}
    # 服务期限
    for pat in [r'(?:服务期限|合同期限|服务周期)[：:]\s*([^\n]{5,80})',
                r'工期[：:]\s*([^\n]{5,80})']:
        m = re.search(pat, text[:5000])
        if m: info['服务期限'] = clean(m.group(1)); break
    # 服务内容
    sec = find_section(text, ['服务内容', '服务要求', '服务范围', '项目内容', '工作内容'])
    if sec:
        info['服务内容'] = clean(sec[:300])
    return info


# ============================================================
#  主提取函数
# ============================================================
def extract_contract(text: str, filename: str) -> Dict:
    """从合同文本提取完整结构化数据"""
    basic = extract_basic(text, filename)
    ptype = basic['项目类型']

    result = {
        '基本信息': {
            '合同编号': basic['合同编号'],
            '财务编号': None,
            '项目名称': basic['项目名称'],
            '甲方': basic['甲方'],
            '乙方': basic['乙方'],
            '项目负责人': None,
            '签订日期': basic['签订日期'],
            '有效期限': basic['有效期限'],
        },
        '项目类型': ptype,
        '合同金额': basic['合同金额'] or {},
        '识别元数据': {
            '识别时间': '2026-07-01T14:00:00+08:00',
            '识别工具': 'contract_extractor_v3',
            '识别状态': '自动提取',
            '备注': f'文本长度{len(text)}字',
        },
        '_raw_text_length': len(text),
    }

    # 研究阶段 (科研类)
    if ptype in ('科研类', '未知'):
        stages = extract_stages(text)
        if stages:
            result['研究阶段'] = stages
            result['识别元数据']['备注'] += f', 提取阶段{len(stages)}个'

    # 服务信息 (服务类)
    if ptype == '服务类':
        sinfo = extract_service(text)
        if sinfo:
            result['服务信息'] = sinfo

    # 付款计划 (always try)
    payments = extract_payments(text)
    if payments:
        result['付款计划'] = payments
        result['识别元数据']['备注'] += f', 提取付款{len(payments)}笔'
    else:
        result['识别元数据']['备注'] += ', 无明确付款计划'

    # 交付物
    deliverables = extract_deliverables(text)
    if deliverables:
        result['交付物'] = deliverables
        result['识别元数据']['备注'] += f', 交付物{len(deliverables)}项'

    return result


# ============================================================
#  批量处理
# ============================================================
def batch_process(ocr_dir: Path, force: bool = False):
    """批量处理OCR合同目录中的所有docx文件"""
    existing = {}
    for f in CACHE_DIR.glob('*标准数据.json'):
        try:
            d = json.load(open(f))
            st = d.get('识别元数据', {}).get('识别状态', '')
            existing[f.stem] = st
        except:
            pass

    docx_files = sorted(ocr_dir.glob('*.docx'))
    print(f'OCR合同库: {len(docx_files)} 个文件')
    print(f'现有标准数据: {len(existing)} 个')

    stats = {'skip': 0, 'process': 0, 'fail': 0}

    for fpath in docx_files:
        fname = fpath.name

        # Check if already processed with "完整" status
        skip = False
        for ek, st in existing.items():
            if (fname[:20] in ek or ek in fname) and st in ('完整', '手动提取'):
                skip = True
                break
        if skip and not force:
            print(f'  ⏭️  {fname[:50]}... (已存在)')
            stats['skip'] += 1
            continue

        print(f'  📄 {fname[:60]}...', end=' ')
        try:
            from docx import Document
            doc = Document(str(fpath))
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            if len(text) < 100:
                print('⚠️ 文本过短')
                stats['fail'] += 1
                continue

            result = extract_contract(text, fname)

            # Determine output filename
            pid = result['基本信息']['合同编号']
            if pid:
                out_name = f'{pid}-标准数据.json'
            else:
                out_name = fname.replace('.docx', '').replace(' ', '')[:40] + '-标准数据.json'

            out_path = CACHE_DIR / out_name
            with open(out_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            print(f'✅ ({len(text)}字)')
            stats['process'] += 1
        except Exception as e:
            print(f'❌ {e}')
            stats['fail'] += 1

    print(f'\n统计: 跳过={stats["skip"]}, 处理={stats["process"]}, 失败={stats["fail"]}')


# ============================================================
#  入口
# ============================================================
if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法:')
        print('  python3 contract_extractor_v3.py <合同.docx>         # 单个合同')
        print('  python3 contract_extractor_v3.py --batch [--force]  # 批量处理')
        print('  python3 contract_extractor_v3.py --status           # 查看状态')
        sys.exit(0)

    if sys.argv[1] == '--batch':
        force = '--force' in sys.argv
        batch_process(OCR_LIB, force)
    elif sys.argv[1] == '--status':
        existing = sorted(CACHE_DIR.glob('*标准数据.json'))
        print(f'标准数据文件: {len(existing)} 个')
        for f in existing:
            try:
                d = json.load(open(f))
                st = d.get('识别元数据', {}).get('识别状态', '?')
                pid = d.get('基本信息', {}).get('合同编号', '')[:25]
                pname = d.get('基本信息', {}).get('项目名称', '')[:40]
                amt = d.get('合同金额', {}).get('含税金额', '')
                stages = len(d.get('研究阶段', []))
                pays = len(d.get('付款计划', []))
                print(f'  [{st:8s}] {pid:25s} | {pname:40s} | ¥{amt or "?":>10s} | S{stages} P{pays}')
            except:
                print(f'  [ERROR] {f.name}')
    else:
        # Single file
        try:
            from docx import Document
            doc = Document(sys.argv[1])
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            result = extract_contract(text, Path(sys.argv[1]).name)
            print(json.dumps(result, ensure_ascii=False, indent=2))
        except Exception as e:
            print(f'Error: {e}')
